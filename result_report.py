from docx import Document
from docx.shared import Inches
import os
from PIL import Image

doc = Document()

path = input('Enter the path of journal :').replace('"', '')
journal_path =os.path.normpath(path)

folder_path = journal_path.split('\\')[:-1]
folder_path = '\\'.join(folder_path)

file_name = journal_path.split('/')[-1].replace('journal', 'results')
file_path = os.path.join(folder_path, file_name)
print('folder', folder_path)

doc.save(file_path)
current_directory = os.getcwd()

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

document_text = extract_text_from_docx(journal_path)
doc.add_heading('Information about the experiment', level=1)
doc.add_paragraph(f'{document_text}')
doc.save(file_path)

def get_images_from_folder(folder_path):
    image_files = []
    for file_name in os.listdir(folder_path):
        if file_name.endswith(('png', 'jpg', 'jpeg', 'bmp', 'gif')):
            image_files.append(os.path.join(folder_path, file_name))
    return image_files

images = get_images_from_folder(folder_path)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'conditions'
hdr_cells[1].text = 'name'
hdr_cells[2].text = 'results'

for number, image_path in enumerate(images):
    file_name, extension = os.path.splitext(image_path)
    file_name = file_name.split('\\')[-1]
    row_cells = table.add_row().cells
    row_cells[0].text = f'{number+1}'
    row_cells[1].text = f'{file_name}'
    row_cells[2].add_paragraph().add_run().add_picture(image_path, width=Inches(1.25))

doc.save(file_path)
print(f"Result report file was saved in folder: {file_path}")